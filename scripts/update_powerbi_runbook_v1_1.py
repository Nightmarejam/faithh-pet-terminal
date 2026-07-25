#!/usr/bin/env python3
"""
Patch TomCatSound_PowerBI_Runbook_Complete.docx → *_v1.1.docx with 2026-04-04 session notes.

Run once against the original Complete.docx only (not v1.1) — re-running would duplicate callouts.

Requires: python-docx (e.g. python3 -m venv /tmp/docx_venv && /tmp/docx_venv/bin/pip install python-docx).
"""
from __future__ import annotations

import sys
from pathlib import Path

from docx import Document
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph


def insert_paragraph_after(paragraph: Paragraph, text: str) -> Paragraph:
    new_p = OxmlElement("w:p")
    paragraph._element.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if text:
        new_para.add_run(text)
    return new_para


def find_paragraph(doc: Document, needle: str) -> Paragraph:
    for p in doc.paragraphs:
        if needle in p.text:
            return p
    raise ValueError(f"Paragraph not found containing: {needle!r}")


def main() -> int:
    src = Path("/mnt/c/Users/jonat/Downloads/TomCatSound_PowerBI_Runbook_Complete.docx")
    dst = Path("/mnt/c/Users/jonat/Downloads/TomCatSound_PowerBI_Runbook_Complete_v1.1.docx")
    if not src.is_file():
        print(f"Missing source: {src}", file=sys.stderr)
        return 1

    doc = Document(str(src))

    # --- Footer / version ---
    footer_needle = "Tom Cat Sound LLC  /  Floating Garden Soundworks  —  Power BI Runbook  —"
    for p in doc.paragraphs:
        if footer_needle in p.text:
            p.text = (
                "Tom Cat Sound LLC  /  Floating Garden Soundworks  —  Power BI Runbook  —  "
                "v1.1 — Updated 2026-04-04 with live session notes"
            )
            break

    # --- Section 4: DAX clarification ---
    p_dax = find_paragraph(doc, "Modeling tab → New measure")
    insert_paragraph_after(
        p_dax,
        "CLARIFICATION: When you click New measure, a formula bar appears at the TOP of the canvas — "
        "type your DAX formula there and press Enter. The Properties panel that opens on the right is "
        "for formatting an existing measure after it is created, not for writing the formula. If you see "
        "a properties panel but no formula bar, make sure you have the Measures table selected in the "
        "Data pane first.",
    )

    # --- Section 5 KPI cards ---
    p_card = find_paragraph(doc, "click the Card visual icon")
    p_kpi_ts = insert_paragraph_after(
        p_card,
        "TROUBLESHOOTING: If you drag a measure directly onto the canvas, Power BI may default to a bar "
        "chart instead of a Card visual. To fix this: click the visual that appeared, then click the "
        "Card icon in the Visualizations pane to convert it. Your data will stay — only the visual type "
        "changes.",
    )
    insert_paragraph_after(
        p_kpi_ts,
        "For uniform card sizing: click one card → Format visual pane → General → Size → note the Width "
        "and Height values → manually set the same values on each other card. This is faster than dragging "
        "to match by eye.",
    )

    # --- Bar chart: replace instruction + notes ---
    p_bar = find_paragraph(doc, "drag Category to Y-axis")
    p_bar.text = (
        "Visualizations pane → Clustered bar chart icon → drag Category from the financials table to Y-axis, "
        "and Amount from the financials table to X-axis. Both fields must come from the financials table "
        "specifically — not from the measures table."
    )
    p_note = insert_paragraph_after(
        p_bar,
        'NOTE: Power BI will auto-name this visual "Sum of Amount by Category". To rename it: click the chart '
        "→ Format visual pane → Title → clear the auto-generated text and type \"2024 P&L Breakdown\". "
        "Power BI auto-names all visuals based on the fields used — always rename titles manually after "
        "adding a visual.",
    )
    insert_paragraph_after(
        p_note,
        "TROUBLESHOOTING: If the chart appears as a filled area or mountain shape instead of horizontal bars, "
        "the wrong visual type was selected. Click the visual, then in the Visualizations pane click the "
        "Clustered bar chart icon (horizontal bars, not vertical columns). The icon looks like three horizontal "
        "bars of different lengths.",
    )

    # --- Gauge ---
    p_gauge = find_paragraph(doc, "The needle should sit at about 8.5%")
    p_fmt = insert_paragraph_after(
        p_gauge,
        "FORMATTING: After adding the gauge, the value will display as a decimal (e.g. 0.09) by default. To show "
        "it as a percentage: Click the gauge → Format visual pane → Callout value → Value decimal places: 1. "
        "Then under Gauge axis → Min value: 0, Max value: 1, and set the format to Percentage. The gauge will "
        'then display "8.52%" correctly.',
    )
    insert_paragraph_after(
        p_fmt,
        "NOTE: The gauge needle sits very low at 8.52% because the scale runs 0–100%. This is intentional — it "
        "visually communicates how thin the margin is relative to a theoretical 100% margin. This is the "
        "correct business interpretation.",
    )

    # --- Section 7: table + heading before "What to Build Next" ---
    ref = find_paragraph(doc, "What to Build Next")
    ref_el = ref._element

    def pre(text: str) -> None:
        new_p = OxmlElement("w:p")
        ref_el.addprevious(new_p)
        Paragraph(new_p, ref._parent).add_run(text)

    table_rows = [
        ["Issue", "Cause", "Fix"],
        [
            "Dragging a measure shows a bar chart, not a card",
            "Power BI defaults to bar chart for numeric fields",
            "Click the visual, then click Card icon in Visualizations pane",
        ],
        [
            "Formula bar doesn't appear when clicking New measure",
            "Measures table not selected",
            "Click the Measures table in the Data pane first, then Modeling → New measure",
        ],
        [
            'Visual title says "Sum of [field] by [field]"',
            "Auto-generated title",
            "Format visual → Title → type your own title",
        ],
        [
            "Bar chart shows as mountain/area shape",
            "Wrong visual type selected",
            "Click visual → Clustered bar chart icon in Visualizations pane",
        ],
        [
            "Gauge shows 0.09 instead of 8.52%",
            "Default format is decimal",
            "Format visual → Gauge axis → Format → Percentage",
        ],
        [
            "Category filter not showing all options",
            "Filter not applied to correct visual",
            "Make sure the visual is selected (blue border) before adding filter",
        ],
        [
            "Ownership shows as 0.34 instead of 34%",
            "Default format is decimal",
            "Click the Ownership column in members table → Column tools → Format → Percentage",
        ],
        [
            "K-1 loss shows as positive number",
            "Amount entered without negative sign",
            "Go to Transform data → find k1_allocations → verify LossShare values are negative",
        ],
    ]

    t = doc.add_table(rows=len(table_rows), cols=3)
    for i, row in enumerate(table_rows):
        for j, cell_text in enumerate(row):
            t.rows[i].cells[j].text = cell_text

    tbl_el = t._tbl
    body = doc.element.body
    body.remove(tbl_el)

    # Before ref (bottom-up): spacer under WBTN, table, spacer, title, top spacer
    pre("")
    ref_el.addprevious(tbl_el)
    pre("")
    pre("Section 7 — Troubleshooting Quick Reference")
    pre("")

    doc.save(str(dst))
    print(f"Wrote {dst}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
