#!/usr/bin/env python3
from __future__ import annotations

from pathlib import Path

from openpyxl import load_workbook
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import RGBColor, Pt
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph

DOWNLOADS = Path('/mnt/c/Users/jonat/Downloads')

BOM_SRC = DOWNLOADS / 'Compressor_BOM_v2.xlsx'
BOM_OUT = DOWNLOADS / 'Compressor bom v3.xlsx'

PLUGIN_OUT = DOWNLOADS / 'Plugin_Development_Runbook.docx'
TRACKER_PATH = DOWNLOADS / 'Gate_Tracking_System.docx'


def insert_paragraph_after(paragraph: Paragraph, text: str = '') -> Paragraph:
    new_p = OxmlElement('w:p')
    paragraph._element.addnext(new_p)
    new_para = Paragraph(new_p, paragraph._parent)
    if text:
        new_para.add_run(text)
    return new_para


def find_paragraph_contains(doc: Document, needle: str) -> Paragraph | None:
    for p in doc.paragraphs:
        if needle in p.text:
            return p
    return None


def update_bom() -> None:
    wb = load_workbook(BOM_SRC)

    # --- Parts List updates ---
    ws = wb['Parts List']

    total_row = None
    for r in range(1, ws.max_row + 1):
        if ws.cell(r, 1).value == 'ESTIMATED TOTAL BUILD COST':
            total_row = r
            break
    if total_row is None:
        raise RuntimeError('Could not find ESTIMATED TOTAL BUILD COST row in Parts List')

    insert_rows = [
        # spacer
        [None, None, None, None, None, None, None, None, None, None],
        # MS ENCODING
        ['MS ENCODING', None, None, None, None, None, None, None, None, None],
        ['MS1', 'Omron G6K-2F-Y relay (x4)', 'DPDT signal relay for MS switching', 'Mouser', 4, 3.50, None, 'P2', 'Not ordered', 'C7'],
        ['MS2', '604 ohm 0.1% resistor (x8)', 'MS encoder/decoder matrix resistors', 'Mouser', 8, 0.30, None, 'P2', 'Not ordered', 'C7'],
        ['MS3', '2-position toggle switch', 'Front panel MS/Stereo mode select', 'Mouser', 1, 4.50, None, 'P2', 'Not ordered', 'C7'],
        ['MS4', 'PCB terminal block 3-pin (x4)', 'Relay connection points', 'Mouser', 4, 1.20, None, 'P2', 'Not ordered', 'C7'],
        ['Subtotal — MS ENCODING', None, 'MS section total: ~$22', None, None, None, None, None, None, None],
        [None, 'NOTE: MS encoder is a passive resistor matrix — two 604 ohm resistors sum L+R for Mid, two 604 ohm resistors diff L-R for Side. Relays switch between L/R stereo mode and M/S mode. Both channels compress independently in either mode.', None, None, None, None, None, None, None, None],
        [None, None, None, None, None, None, None, None, None, None],
        # SOFT CLIPPER
        ['SOFT CLIPPER', None, None, None, None, None, None, None, None, None],
        ['SCP1', '12AX7 / ECC83 tube', 'Soft clipper tube (half triode used)', 'Tube Depot', 1, 12.00, None, 'P2', 'Not ordered', 'C8'],
        ['SCP2', '9-pin noval tube socket', 'PCB mount', 'Mouser', 1, 3.50, None, 'P2', 'Not ordered', 'C8'],
        ['SCP3', '100k audio taper pot', 'Soft clip drive control', 'Mouser', 1, 3.50, None, 'P2', 'Not ordered', 'C8'],
        ['SCP4', '47k resistor 1%', 'Plate load resistor', 'Mouser', 2, 0.15, None, 'P2', 'Not ordered', 'C8'],
        ['SCP5', '1.5k resistor 1%', 'Cathode resistor', 'Mouser', 2, 0.15, None, 'P2', 'Not ordered', 'C8'],
        ['SCP6', '22uF 50V electrolytic cap', 'Cathode bypass', 'Mouser', 2, 0.40, None, 'P2', 'Not ordered', 'C8'],
        ['SCP7', '0.1uF film cap', 'Coupling cap to output stage', 'Mouser', 2, 0.60, None, 'P2', 'Not ordered', 'C8'],
        ['Subtotal — SOFT CLIPPER', None, 'Soft clipper section total: ~$25', None, None, None, None, None, None, None],
        [None, 'NOTE: Soft clipper placed AFTER NE5532 makeup gain, BEFORE output transformer. Half triode per channel. Drive control sets clip threshold. Transformer further softens any clipping character. Gentle drive setting is transparent — push it for harmonic density. Uses same B+ supply as main tube stages.', None, None, None, None, None, None, None, None],
        [None, None, None, None, None, None, None, None, None, None],
    ]

    ws.insert_rows(total_row, amount=len(insert_rows))

    start = total_row
    for i, row in enumerate(insert_rows):
        rr = start + i
        for c, val in enumerate(row, start=1):
            if val is not None:
                ws.cell(rr, c, val)

    # Ext $ formulas for part lines
    for rr in range(start, start + len(insert_rows)):
        part_id = ws.cell(rr, 1).value
        if isinstance(part_id, str) and part_id.startswith(('MS', 'SCP')):
            ws.cell(rr, 7, f'=E{rr}*F{rr}')

    # Subtotals
    ms_first = start + 2
    ms_last = start + 5
    ws.cell(start + 6, 7, f'=SUM(G{ms_first}:G{ms_last})')

    sc_first = start + 10
    sc_last = start + 16
    ws.cell(start + 17, 7, f'=SUM(G{sc_first}:G{sc_last})')

    # Update grand total row/formula (old row shifted)
    new_total_row = total_row + len(insert_rows)
    ws.cell(new_total_row, 1, 'ESTIMATED TOTAL BUILD COST')
    ws.cell(new_total_row, 7, f'=SUM(G5:G{new_total_row-1})')

    # --- Signal Chain updates ---
    sig = wb['Signal Chain']
    # Insert after existing stage 6 row (current row 10)
    sig.insert_rows(11, amount=2)
    sig.cell(11, 1, 'MS Encoder')
    sig.cell(11, 2, 'MS ENCODER')
    sig.cell(11, 3, 'Splits L/R to M/S before gain reduction')
    sig.cell(11, 4, 'Relay matrix + resistor network')
    sig.cell(11, 5, 'Switchable — front panel toggle')

    sig.cell(12, 1, 'Soft Clipper')
    sig.cell(12, 2, 'SOFT CLIPPER')
    sig.cell(12, 3, 'Rounds transients post makeup gain')
    sig.cell(12, 4, '12AX7 half triode per channel')
    sig.cell(12, 5, 'Drive control on front panel')

    # --- Panel Layout updates ---
    panel = wb['Panel Layout Reference']
    # Insert around center-control rows (after Stereo Link)
    panel.insert_rows(16, amount=2)
    panel.cell(16, 1, 'MS/STEREO toggle switch')
    panel.cell(16, 2, '1')
    panel.cell(16, 3, '6mm / DPDT')
    panel.cell(16, 4, 'Center zone, below Link toggle')
    panel.cell(16, 5, 'Carling DPDT')

    panel.cell(17, 1, 'CLIP DRIVE knob (100k pot)')
    panel.cell(17, 2, '2')
    panel.cell(17, 3, '6mm shaft / 9mm')
    panel.cell(17, 4, 'Right of makeup gain, left channel and right channel rows')
    panel.cell(17, 5, 'Alpha audio taper 100k')

    # Add capacity note before rear panel note
    rear_row = None
    for r in range(1, panel.max_row + 1):
        v = panel.cell(r, 1).value
        if isinstance(v, str) and v.startswith('REAR PANEL:'):
            rear_row = r
            break
    if rear_row is None:
        rear_row = panel.max_row + 1
    panel.insert_rows(rear_row, amount=1)
    panel.cell(
        rear_row,
        1,
        'NOTE: Front panel now has 2 additional controls per channel (clip drive) plus 1 center toggle (MS mode). 2U panel is at capacity — no further controls should be added without moving to 3U chassis.',
    )

    # --- Order Priority updates ---
    if 'Order Priority' in wb.sheetnames:
        op = wb['Order Priority']
        row = op.max_row + 2
    else:
        op = wb.create_sheet('Order Priority')
        op.cell(1, 1, 'ORDER PRIORITY')
        op.cell(3, 1, 'Phase')
        op.cell(3, 2, 'Item')
        op.cell(3, 3, 'Est Cost')
        op.cell(3, 4, 'Notes')
        row = 4

    op.cell(row, 1, 'Phase 2 addition')
    op.cell(row, 2, 'MS Encoding components')
    op.cell(row, 3, 22)
    op.cell(row + 1, 1, 'Phase 2 addition')
    op.cell(row + 1, 2, 'Soft Clipper components')
    op.cell(row + 1, 3, 25)
    op.cell(row + 2, 1, 'Phase 2 revised total')
    op.cell(row + 2, 3, '=C{}+C{}'.format(row, row + 1))
    op.cell(row + 2, 4, 'Add ~$47 to existing Phase 2 figure')

    wb.save(BOM_OUT)


def add_blue_heading(doc: Document, text: str, level: int = 2) -> None:
    p = doc.add_paragraph(text)
    p.style = f'Heading {level}'
    for run in p.runs:
        run.font.color.rgb = RGBColor(31, 78, 121)


def add_code_block(doc: Document, text: str) -> None:
    p = doc.add_paragraph(text)
    p.style = doc.styles['No Spacing']
    for run in p.runs:
        run.font.name = 'Consolas'
        run.font.size = Pt(10)


def create_plugin_runbook() -> None:
    doc = Document()

    title = doc.add_paragraph('FGS Audio — Plugin Development Runbook')
    title.style = 'Title'
    subtitle = doc.add_paragraph('JUCE Setup and First Plugin Guide')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph('Version: 1.0 — 2026-04-04')
    doc.add_paragraph('For: Development partner setup')

    add_blue_heading(doc, 'Section 1 — Overview', 2)
    for line in [
        'JUCE is a C++ framework for building audio plugins (VST3, AU, AAX).',
        'Cross platform — build once, deploy to Windows and macOS.',
        'Industry standard — used by Waves, iZotope, Native Instruments.',
        'Free for open source / personal use, paid license for commercial release.',
        'Our first plugin: a transformer color box (harmonic saturation) modeled after the output stage of the FGS mastering compressor.',
    ]:
        doc.add_paragraph(line, style='List Bullet')

    add_blue_heading(doc, 'Section 2 — Prerequisites', 2)
    doc.add_paragraph('Windows:', style='List Bullet')
    for line in [
        'Visual Studio 2022 Community (free) — https://visualstudio.microsoft.com (install "Desktop development with C++").',
        'Git — https://git-scm.com/download/win',
        'CMake 3.22+ — https://cmake.org/download/',
    ]:
        doc.add_paragraph(line, style='List Bullet 2')
    doc.add_paragraph('macOS (if applicable):', style='List Bullet')
    for line in [
        'Xcode (latest) from Mac App Store.',
        'Xcode Command Line Tools: run `xcode-select --install` in Terminal.',
        'Git (included with Xcode CLT).',
        'CMake: `brew install cmake` (requires Homebrew).',
    ]:
        doc.add_paragraph(line, style='List Bullet 2')

    add_blue_heading(doc, 'Section 3 — Install JUCE', 2)
    steps = [
        'Go to https://juce.com/get-juce/',
        'Click "Download JUCE" and select the latest stable release.',
        'Extract to a permanent location: Windows `C:\\JUCE\\` / macOS `~/JUCE/` (do not leave in Downloads).',
        'Open Projucer inside the JUCE folder.',
        'On first launch, select "Personal / GPL" for now.',
    ]
    for s in steps:
        doc.add_paragraph(s, style='List Number')

    add_blue_heading(doc, 'Section 4 — Clone the FGS Plugin Repository', 2)
    doc.add_paragraph('NOTE: Jonathan creates the repository and sends collaborator invite before clone.')
    add_code_block(doc, '# Windows (Git Bash or Command Prompt):\ncd C:\\Users\\[username]\\Documents\ngit clone https://github.com/[repo-to-be-created]/fgs-plugins.git\ncd fgs-plugins')
    add_code_block(doc, '# macOS:\ncd ~/Documents\ngit clone https://github.com/[repo-to-be-created]/fgs-plugins.git\ncd fgs-plugins')

    add_blue_heading(doc, 'Section 5 — Open and Build the First Project', 2)
    doc.add_paragraph('In Projucer:', style='List Bullet')
    for s in [
        'File → Open → `transformer-color-box/TransformerColorBox.jucer`',
        'File → Save Project and Open in IDE',
        'IDE opens in Visual Studio (Windows) or Xcode (macOS)',
    ]:
        doc.add_paragraph(s, style='List Bullet 2')
    doc.add_paragraph('Visual Studio (Windows):', style='List Bullet')
    for s in [
        'Select `TransformerColorBox_VST3` target',
        'Set configuration to `Release`',
        'Build → Build Solution (Ctrl+Shift+B)',
        'Output: `C:\\Users\\[username]\\AppData\\Roaming\\VST3\\TransformerColorBox.vst3`',
    ]:
        doc.add_paragraph(s, style='List Bullet 2')
    doc.add_paragraph('Xcode (macOS):', style='List Bullet')
    for s in [
        'Select `TransformerColorBox - AU` scheme',
        'Product → Build (Cmd+B)',
        'Output: `~/Library/Audio/Plug-Ins/Components/TransformerColorBox.component`',
    ]:
        doc.add_paragraph(s, style='List Bullet 2')

    add_blue_heading(doc, 'Section 6 — Test in a DAW', 2)
    doc.add_paragraph('Recommended DAW: Reaper (https://www.reaper.fm/download.php)', style='List Bullet')
    for s in [
        'Open Reaper',
        'Options → Preferences → Plug-ins → VST',
        'Add VST path: `C:\\Users\\[username]\\AppData\\Roaming\\VST3\\`',
        'Click Re-scan',
        'Create track → FX → search "Transformer" → load plugin',
    ]:
        doc.add_paragraph(s, style='List Number')
    for s in [
        'Input: audio signal',
        'Drive 0-100% increases harmonic saturation',
        'Output control for level compensation',
        'Character warm/subtle at low drive, more color at high drive',
        'Bypass for A/B comparison',
    ]:
        doc.add_paragraph(s, style='List Bullet')

    add_blue_heading(doc, 'Section 7 — Workflow for New Plugins', 2)
    for s in [
        'Jonathan creates a gate document describing plugin behavior.',
        'Partner creates new Projucer project from `PluginTemplate` starter.',
        'DSP code in `PluginProcessor.cpp`.',
        'UI code in `PluginEditor.cpp`.',
        'Both commit to shared GitHub repository.',
        'Jonathan tests in Windows + FAITHH environment.',
        'Partner tests on their DAW setup.',
        'Gate closes when both systems produce identical output.',
    ]:
        doc.add_paragraph(s, style='List Number')

    add_blue_heading(doc, 'Section 8 — Communication and Handoff', 2)
    for s in [
        'All plugin specs are documented as gate files in repo.',
        'Questions go in GitHub Issues.',
        'Each plugin gets its own branch: `plugin/transformer-color-box`.',
        'Merge to main only when gate criteria are met.',
    ]:
        doc.add_paragraph(s, style='List Bullet')

    add_blue_heading(doc, 'Section 9 — Troubleshooting', 2)
    table = doc.add_table(rows=1, cols=3)
    hdr = table.rows[0].cells
    hdr[0].text = 'Problem'
    hdr[1].text = 'Likely Cause'
    hdr[2].text = 'Fix'
    rows = [
        ("Plugin doesn't appear in DAW", 'Wrong VST path or no rescan', 'Re-scan VST paths in DAW preferences'),
        ('Build fails: JUCE not found', 'Projucer path wrong', 'Set JUCE modules path to C:\\JUCE\\modules in Projucer Preferences'),
        ('Visual Studio not found', 'Missing install/workload', 'Install VS2022 with C++ Desktop workload'),
        ('Plugin crashes DAW', 'Debug build loaded', 'Rebuild in Release configuration'),
        ('macOS plugin not authorized', 'Gatekeeper block', 'System Preferences → Security → Allow'),
        ('Git push rejected', 'Not added as collaborator', 'Ask Jonathan to send GitHub invitation'),
    ]
    for r in rows:
        c = table.add_row().cells
        c[0].text, c[1].text, c[2].text = r

    add_blue_heading(doc, 'Appendix A — Recommended Learning Resources', 2)
    for s in [
        'The Audio Programmer (YouTube): https://www.youtube.com/@TheAudioProgrammer',
        'JUCE documentation: https://docs.juce.com',
        'ADC talks (YouTube)',
        'Will Pirkle — Designing Audio Effect Plugins in C++',
    ]:
        doc.add_paragraph(s, style='List Bullet')

    add_blue_heading(doc, 'Appendix B — First Plugin Spec: Transformer Color Box', 2)
    doc.add_paragraph('Design intent: Model the harmonic character of the Cinemag CMOB-2H output transformer in the FGS mastering compressor.')
    doc.add_paragraph('Controls:', style='List Bullet')
    for s in [
        'DRIVE (0-100%): increases even-order harmonic content',
        'IRON / NICKEL switch: iron=warm, nickel=clean',
        'OUTPUT (-12 to +6 dB): level compensation',
        'BYPASS: true bypass comparison',
    ]:
        doc.add_paragraph(s, style='List Bullet 2')
    doc.add_paragraph('DSP approach:', style='List Bullet')
    for s in [
        'Even harmonic saturation via soft clip waveshaper emphasizing 2nd and 4th harmonics',
        'Iron mode: gentle LF shelving + stronger 2nd harmonic',
        'Nickel mode: flatter response + stronger 4th harmonic',
        'Oversampling: 4x minimum to reduce aliasing',
    ]:
        doc.add_paragraph(s, style='List Bullet 2')
    doc.add_paragraph('Gate criteria:', style='List Bullet')
    for s in [
        'Spectrum shows expected harmonic profile at 50% drive',
        'Iron and Nickel modes show measurable spectral difference',
        'No aliasing artifacts above 18kHz at 44.1kHz',
        'Passes null test against bypass within 0.1 dB',
    ]:
        doc.add_paragraph(s, style='List Bullet 2')

    # Footer
    footer = doc.sections[0].footer.paragraphs[0]
    footer.text = 'FGS Audio — Plugin Development Runbook — v1.0 — 2026-04-04'
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.save(PLUGIN_OUT)


def update_tracker_doc() -> None:
    doc = Document(TRACKER_PATH)

    # Add Track 5 references near top track line
    top_tracks = find_paragraph_contains(doc, 'Track 1: Resonant Compressor Build')
    if top_tracks is not None:
        p1 = insert_paragraph_after(top_tracks, 'Track 5: Plugin Development')
        p2 = insert_paragraph_after(p1, 'Partner runbook: Plugin_Development_Runbook.docx (in Downloads folder)')
        p3 = insert_paragraph_after(p2, 'GitHub repo: to be created by Jonathan — see PL1 gate criteria')
        insert_paragraph_after(p3, 'First plugin: Transformer Color Box — spec in Appendix B of runbook')

    # Add compressor notes under TRACK 1 section heading
    track1_heading = find_paragraph_contains(doc, 'TRACK 1: RESONANT COMPRESSOR BUILD')
    if track1_heading is not None:
        t1 = insert_paragraph_after(track1_heading, 'BOM v3 adds MS encoding (Gate C7) and soft clipper (Gate C8)')
        t2 = insert_paragraph_after(t1, 'See Compressor bom v3.xlsx for updated parts and pricing')
        insert_paragraph_after(t2, 'Front panel at capacity with these additions — 3U if further controls needed')

    doc.save(TRACKER_PATH)


def main() -> int:
    update_bom()
    create_plugin_runbook()
    update_tracker_doc()
    print('Updated files:')
    print(BOM_OUT)
    print(PLUGIN_OUT)
    print(TRACKER_PATH)
    return 0


if __name__ == '__main__':
    raise SystemExit(main())
